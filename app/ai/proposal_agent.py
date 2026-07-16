import io
import json
import logging
import re
from datetime import datetime
from uuid import UUID

import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from python_docx_replace import docx_replace  # Installed utility to prevent broken run fragments
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.lead import Lead
from app.models.team import Team
from app.models.proposal import ProposalTemplate  # Import the ProposalTemplate model

logger = logging.getLogger(__name__)


class ProposalAgent:
    """
    Dedicated agent responsible for generating tailored sales proposals 
    for B2B prospects. It analyzes context and style guides, retrieves templates, 
    and outputs perfectly formatted, user-colored Word documents.
    """

    BASE_URL = "https://api.groq.com/openai/v1"
    MODEL = "openai/gpt-oss-120b"

    def __init__(self, db_session: AsyncSession):
        self.db = db_session

    async def run(
        self, 
        user_prompt: str, 
        team_id: UUID, 
        lead_id: UUID | None = None, 
        template_id: UUID | None = None
    ) -> tuple[str, io.BytesIO]:
        """
        Primary entry point. Resolves lead details and the optional custom template,
        queries the LLM for proposal content AND style preferences,
        and produces a tailored .docx stream.
        """
        # 1. Fetch Company/Team Context (ICP Context)
        team_query = select(Team).where(Team.id == team_id)
        team_res = await self.db.execute(team_query)
        team = team_res.scalar_one_or_none()
        icp_context = team.icp if team and team.icp else "Standard Enterprise B2B SaaS and Professional Solutions."

        # 2. Fetch Template File if specified
        template_bytes = None
        if template_id:
            template_query = select(ProposalTemplate).where(ProposalTemplate.id == template_id)
            template_res = await self.db.execute(template_query)
            template_record = template_res.scalar_one_or_none()
            
            if template_record and template_record.file_url:
                try:
                    # Download the template .docx file from your storage (S3/Cloud storage)
                    async with httpx.AsyncClient() as client:
                        resp = await client.get(template_record.file_url)
                        resp.raise_for_status()
                        template_bytes = io.BytesIO(resp.content)
                except Exception as e:
                    logger.warning(f"Failed to fetch template from S3: {e}. Falling back to clean generation.")

        # 3. Fetch Target Lead context if provided
        lead_details = ""
        if lead_id:
            lead_query = select(Lead).where(Lead.id == lead_id)
            lead_res = await self.db.execute(lead_query)
            lead = lead_res.scalar_one_or_none()
            if lead:
                lead_details = (
                    f"Target Prospect: {lead.name}\n"
                    f"Company: {lead.company_name or 'N/A'}\n"
                    f"Job Title: {lead.job_title or 'N/A'}\n"
                    f"Prospect Email: {lead.email or 'N/A'}\n"
                )

        combined_prompt = f"{lead_details}\nUser Specific Goals & Styling preferences: {user_prompt}" if lead_details else user_prompt

        # 4. Call Groq to generate copy and styles (colors)
        try:
            proposal_title, proposal_data, style_config = await self._compile_proposal_data(
                user_prompt=combined_prompt, 
                icp_context=icp_context,
                has_template=bool(template_bytes)
            )
        except Exception as e:
            logger.error(f"ProposalAgent text compilation failed: {e}", exc_info=True)
            raise RuntimeError("ProposalAgent was unable to successfully compile the proposal copy.")

        # 5. Generate Styled/Templated DOCX binary stream
        try:
            doc_stream = self.create_proposal_document(
                proposal_title=proposal_title, 
                proposal_data=proposal_data,
                style_config=style_config,
                template_bytes=template_bytes
            )
            return f"{proposal_title}.docx", doc_stream
        except Exception as e:
            logger.error(f"DOCX compilation failed: {e}", exc_info=True)
            raise RuntimeError("ProposalAgent failed to construct the Microsoft Word file format.")

    async def _compile_proposal_data(self, user_prompt: str, icp_context: str, has_template: bool) -> tuple[str, dict, dict]:
        """
        Queries LLM to return structured proposal data, document metadata, 
        and color styling values extracted directly from the user's input.
        """
        system_prompt = f"""You are an elite enterprise B2B sales strategist. 
        Take the user's prompt request and draft a highly persuasive, detailed business proposal.
        Assume the current context year is 2026.
        There should be no mention that it was written by AI. Do not assume facts outside the user's message.

        Analyze the user's instructions for design color requests (e.g. "use bright blue", "dark mode", "forest green", "hex #E06666").
        Extract or infer a primary, secondary, and text color hex code for formatting. 
        If nothing is specified, fallback to professional Navy ('#1B365D'), Slate Blue ('#4A777A'), and Dark Charcoal ('#333333').

        Our Company Context / Profile:
        \"\"\"
        {icp_context}
        \"\"\"

        Return ONLY a valid JSON object matching this exact structure:
        {{
            "document_title": "A short clean file-safe title (e.g., SalesSync_Acme_Growth_Proposal)",
            "styles": {{
                "primary_color_hex": "HEX code starting with # (e.g., #1B365D)",
                "secondary_color_hex": "HEX code starting with # (e.g., #4B6B94)",
                "text_color_hex": "HEX code starting with # (e.g., #333333)"
            }},
            "proposal_data": {{
                "executive_summary": "Deep, highly persuasive paragraph highlighting key business drivers.",
                "problem_statement": "A technical breakdown of client operational pain points.",
                "proposed_solution": "A detailed step-by-step resolution strategy.",
                "investment_and_pricing": "Clear commercial pricing packages and timeline."
            }}
        }}
        """
        
        payload = {
            "model": self.MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.5,
            "max_tokens": 2000,
            "response_format": {"type": "json_object"}
        }

        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(
                f"{self.BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {settings.grok_api_key}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )
            response.raise_for_status()
            data = response.json()
            content = data["choices"][0]["message"]["content"].strip()
            
            cleaned = re.sub(r"^```json\s*|\s*```$", "", content, flags=re.IGNORECASE)
            parsed_proposal = json.loads(cleaned)
            
            title = parsed_proposal.get("document_title", "Business_Proposal")
            styles = parsed_proposal.get("styles", {})
            proposal_data = parsed_proposal.get("proposal_data", {})
            
            return title, proposal_data, styles

    def create_proposal_document(
        self, 
        proposal_title: str, 
        proposal_data: dict, 
        style_config: dict, 
        template_bytes: io.BytesIO | None = None
    ) -> io.BytesIO:
        """
        Merges generated content and dynamic colors.
        If a template is provided (e.g., containing just a logo/header), 
        it appends the generated content to it. Otherwise, starts from scratch.
        """
        # Convert hex codes into python-docx RGBColor elements
        def hex_to_rgb(hex_str: str) -> RGBColor:
            hex_str = hex_str.lstrip('#')
            if len(hex_str) != 6:
                return RGBColor(0x1B, 0x36, 0x5D)
            return RGBColor(*(int(hex_str[i:i+2], 16) for i in (0, 2, 4)))

        # Extract user's custom colors (or defaults)
        p_color = hex_to_rgb(style_config.get("primary_color_hex", "#1B365D"))
        s_color = hex_to_rgb(style_config.get("secondary_color_hex", "#4A777A"))
        t_color = hex_to_rgb(style_config.get("text_color_hex", "#333333"))

        # 1. Initialize or Load Document
        if template_bytes:
            # Load the user's template (keeps logo, headers, margins, etc.)
            doc = Document(template_bytes)
            
            # If the template has empty placeholder paragraphs at the end, 
            # we can append directly. Let's add a spacing paragraph first.
            doc.add_paragraph()
        else:
            # Fallback to standard blank document
            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

        # 2. Apply base paragraph style colors
        if 'Normal' in doc.styles:
            doc.styles['Normal'].font.name = 'Arial'
            doc.styles['Normal'].font.size = Pt(11)
            doc.styles['Normal'].font.color.rgb = t_color

        # 3. Append Header Title Layout
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(proposal_title.replace("_", " ").upper())
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = p_color  # USER PRIMARY COLOR
        
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run("Strategic Growth Proposal")
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True
        meta_run.font.color.rgb = s_color  # USER SECONDARY COLOR
        
        # Divider Line
        divider_p = doc.add_paragraph()
        divider_run = divider_p.add_run("_" * 60)
        divider_run.font.color.rgb = RGBColor(0xD3, 0xD3, 0xD3)
        doc.add_paragraph() 

        # 4. Map and Write Content Sections
        sections_map = {
            "executive_summary": "1. Executive Summary",
            "problem_statement": "2. Problem Statement & Operational Context",
            "proposed_solution": "3. Strategic Roadmap & Proposed Implementation",
            "investment_and_pricing": "4. Commercial Terms & Financial Scope"
        }

        for key, section_heading in sections_map.items():
            text_block = proposal_data.get(key, "Section contents skipped.")
            
            # Add section heading
            heading_p = doc.add_paragraph()
            heading_run = heading_p.add_run(section_heading)
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            heading_run.font.color.rgb = p_color  # USER PRIMARY COLOR
            heading_p.paragraph_format.space_before = Pt(16)
            heading_p.paragraph_format.space_after = Pt(6)
            
            # Add section body text
            body_p = doc.add_paragraph()
            body_run = body_p.add_run(text_block)
            body_run.font.color.rgb = t_color  # USER TEXT COLOR
            body_p.paragraph_format.space_after = Pt(12)
            body_p.paragraph_format.line_spacing = 1.15

        # Save to byte stream
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        return out_stream